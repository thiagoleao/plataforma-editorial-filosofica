import hmac
import io
import json
import math
import os
import re
import uuid
from datetime import datetime, timezone
from functools import wraps
from urllib.parse import quote

import psycopg2
from docx import Document
from docx.shared import Inches, Pt
from flask import Flask, Response, jsonify, request
from openai import OpenAI
from pgvector.psycopg2 import register_vector
from psycopg2.extras import Json, RealDictCursor

app = Flask(__name__)

INSTANCE_CONNECTION_NAME = os.environ["INSTANCE_CONNECTION_NAME"]
DB_NAME = os.environ.get("DB_NAME", "plataforma_editorial_filosofica")
DB_USER = os.environ.get("DB_USER", "editorial_app")
DB_PASSWORD = os.environ["DB_PASSWORD"]
SERVICE_API_KEY = os.environ["SERVICE_API_KEY"]
OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]

EMBEDDING_MODEL = "text-embedding-3-small"
EMBEDDING_MAX_CHARS = 20000
CHAT_MODEL = "gpt-4.1-mini"
RELATION_TYPES = {"cooccurrence", "causal", "evolutionary", "pre_requisito", "contraste", "manifestacao_de"}

# recalculate_concept_graph() varre editorial.concepts/concept_relations/knowledge_cards quase
# por inteiro a cada chamada — chamadas concorrentes (ex.: upserts em paralelo) tomam locks de
# linha em ordens diferentes e o Postgres derruba uma delas com DeadlockDetected. Serializar via
# advisory lock transacional (libera sozinho no commit/rollback) em vez de mudar a lógica da função.
CONCEPT_GRAPH_LOCK_KEY = 741125

openai_client = OpenAI(api_key=OPENAI_API_KEY)


def get_connection():
    connection = psycopg2.connect(
        dbname=DB_NAME,
        user=DB_USER,
        password=DB_PASSWORD,
        host=f"/cloudsql/{INSTANCE_CONNECTION_NAME}",
        connect_timeout=10,
    )
    register_vector(connection)
    return connection


def generate_embedding(text):
    """Gera embedding via OpenAI (ADR-012 §4). Retorna None em caso de falha —
    a falta de embedding nunca deve bloquear a preservação do conteúdo original."""
    text = (text or "").strip()[:EMBEDDING_MAX_CHARS]
    if not text:
        return None
    try:
        response = openai_client.embeddings.create(model=EMBEDDING_MODEL, input=text)
        return response.data[0].embedding
    except Exception:
        app.logger.exception("Failed to generate embedding")
        return None


SEGMENT_INSIGHT_SYSTEM_PROMPT = """Você é um assistente editorial que lê trechos de canalizações \
filosóficas/psicológicas e produz cartões de contextualização — nunca reescreve, resume ou \
parafraseia o texto original.

Para cada conceito central e genuinamente distinto no trecho (no máximo 3), produza um cartão com:
- concept_title: título curto do conceito
- explanation: explicação do conceito
- philosophical_context: com que linhagem/escola de pensamento ele dialoga
- practical_application: como se aplica na vida prática
- related_concepts: lista curta de conceitos (nomes canônicos em português) associados

No campo philosophical_context, use SEMPRE linguagem de aproximação — "ressoa com", "dialoga com" —
nunca de equivalência ("isto é X", "isto vem de X"). Cite conceitos nomeados específicos (ex. "moral
do rebanho", "má-fé sartriana", "biopoder foucaultiano"), nunca referências vagas de autoridade. A
fonte é canalização espiritual/psicológica, não filosofia acadêmica — trate qualquer conexão como
interpretação editorial, nunca como fato sobre a origem do texto.

Depois, se dois ou mais conceitos (dos cartões acima ou da lista de conceitos já associados ao
segmento, fornecida pelo usuário) se relacionarem de forma clara e evidente no texto, classifique a
relação em exatamente um destes tipos: pre_requisito, contraste, manifestacao_de, causal,
evolutionary. Só inclua relações com evidência clara — não force relações.

Responda em português, em JSON estrito, sem texto fora do JSON, no formato:
{"insights": [{"concept_title": "...", "explanation": "...", "philosophical_context": "...",
"practical_application": "...", "related_concepts": ["..."]}],
"relations": [{"concept_a": "...", "concept_b": "...", "relation_type": "..."}]}"""


def generate_segment_insights(segment_text, existing_concept_names):
    """Gera cartões de insight filosófico + relações tipadas para um segmento via chat
    completion (ADR-022 §2-3). Nunca reescreve o segmento — produz comentário interpretativo
    separado. Retorna ([], []) em caso de falha; a ausência de insight nunca bloqueia nada."""
    user_prompt = (
        f"Conceitos já associados a este segmento (use como referência, se úteis): "
        f"{', '.join(existing_concept_names) if existing_concept_names else 'nenhum'}\n\n"
        f"Trecho:\n{(segment_text or '').strip()[:8000]}"
    )
    try:
        response = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": SEGMENT_INSIGHT_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
        )
        data = json.loads(response.choices[0].message.content)
        return data.get("insights") or [], data.get("relations") or []
    except Exception:
        app.logger.exception("Failed to generate segment insights")
        return [], []


def apply_relation_types(cursor, name_to_id, relations):
    """Aplica tipagem de concept_relations inferida por LLM (ADR-022 §3, completa a
    ADR-012 §2). Cria a relação se ainda não existir, atualiza o tipo se já existir."""
    for relation in relations or []:
        relation_type = relation.get("relation_type")
        if relation_type not in RELATION_TYPES:
            continue
        id_a = name_to_id.get(relation.get("concept_a"))
        id_b = name_to_id.get(relation.get("concept_b"))
        if not id_a or not id_b or id_a == id_b:
            continue
        cursor.execute("""
            INSERT INTO editorial.concept_relations (concept_a_id, concept_b_id, relation_type, cooccurrence_count, last_observed_at)
            VALUES (LEAST(%(a)s::uuid, %(b)s::uuid), GREATEST(%(a)s::uuid, %(b)s::uuid), %(relation_type)s, 1, NOW())
            ON CONFLICT (concept_a_id, concept_b_id) DO UPDATE SET
                relation_type = EXCLUDED.relation_type,
                last_observed_at = NOW()
        """, {"a": id_a, "b": id_b, "relation_type": relation_type})


CHAPTER_SUGGESTION_SYSTEM_PROMPT = """Você é um assistente editorial que sugere título e \
resumo curtos para um capítulo candidato de livro, a partir de um conceito central e uma \
lista de títulos de fontes (segmentos canalizados e fichas) já selecionadas para esse \
capítulo. Você não lê o conteúdo integral das fontes — só os títulos, fornecidos abaixo — \
e nunca gera nem sugere texto para o corpo do capítulo em si, só a etiqueta editorial.

Responda em português, em JSON estrito: {"title": "...", "summary": "..."}
O title é curto (até 8 palavras). O summary é 1-2 frases descrevendo do que trata o \
capítulo candidato, para alguém decidir se vale promover essa sugestão para um projeto real."""


def generate_chapter_suggestion_label(concept_name, source_titles):
    """Gera título + resumo curtos para um capítulo candidato via chat completion
    (ADR-019 §4). Retorna (None, None) em caso de falha — nunca bloqueia o lote inteiro."""
    user_prompt = (
        f"Conceito central: {concept_name}\n\nTítulos das fontes selecionadas:\n"
        + "\n".join(f"- {title}" for title in source_titles)
    )
    try:
        response = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": CHAPTER_SUGGESTION_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
        )
        data = json.loads(response.choices[0].message.content)
        return data.get("title"), data.get("summary")
    except Exception:
        app.logger.exception("Failed to generate chapter suggestion label")
        return None, None


def build_proposed_sources(cursor, concept_ids, limit=20, min_importance=25, similarity_threshold=0.92):
    """Monta uma lista ordenada de fontes candidatas (segmentos + fichas) para um escopo
    de conceitos, com dedup por similaridade de embedding (ADR-013 §3). Reaproveitado por
    /chapters/<id>/propose e pela geração de sugestões de capítulo (ADR-019 §4). Retorna
    uma lista de tuplas (kind, id), kind em ('segment', 'card')."""
    cursor.execute("""
        SELECT DISTINCT cs.id AS segment_id, cs.is_channeled, cs.editorial_relevance, cs.embedding
        FROM editorial.content_segments cs
        JOIN editorial.segment_concepts sc ON sc.segment_id = cs.id
        WHERE sc.concept_id = ANY(%s::uuid[])
        ORDER BY cs.is_channeled DESC, cs.editorial_relevance DESC
        LIMIT %s
    """, (concept_ids, limit * 2))
    segment_candidates = cursor.fetchall()

    cursor.execute("""
        SELECT DISTINCT kc.id AS card_id, kc.importance_score, kc.embedding
        FROM editorial.knowledge_cards kc
        JOIN editorial.card_concepts cc ON cc.card_id = kc.id
        WHERE cc.concept_id = ANY(%s::uuid[]) AND kc.importance_score >= %s
        ORDER BY kc.importance_score DESC
        LIMIT %s
    """, (concept_ids, min_importance, limit * 2))
    card_candidates = cursor.fetchall()

    kept = []
    kept_embeddings = []

    def try_keep(kind, item_id, embedding):
        if len(kept) >= limit:
            return
        if embedding is not None and any(
            cosine_similarity(embedding, other) >= similarity_threshold for other in kept_embeddings
        ):
            return
        kept.append((kind, item_id))
        kept_embeddings.append(embedding)

    for row in segment_candidates:
        try_keep("segment", row["segment_id"], row["embedding"])
    for row in card_candidates:
        try_keep("card", row["card_id"], row["embedding"])

    return kept


def find_chapter_suggestion_candidates(cursor, batch_size):
    """Acha conceitos com boa cobertura de conteúdo mas pouco representados em capítulos
    já existentes (ADR-019 §4) — subqueries escalares evitam o fan-out de fazer dois LEFT
    JOIN (segmentos e fichas) na mesma query."""
    cursor.execute("""
        WITH concept_stats AS (
            SELECT
                c.id, c.canonical_name, c.importance_score,
                (SELECT COUNT(*) FROM editorial.segment_concepts sc WHERE sc.concept_id = c.id) +
                (SELECT COUNT(*) FROM editorial.card_concepts cc WHERE cc.concept_id = c.id) AS content_count,
                (
                    SELECT COUNT(DISTINCT cs.id) FROM editorial.chapter_sources cs
                    WHERE cs.segment_id IN (SELECT segment_id FROM editorial.segment_concepts WHERE concept_id = c.id)
                       OR cs.knowledge_card_id IN (SELECT card_id FROM editorial.card_concepts WHERE concept_id = c.id)
                ) AS chapter_source_count
            FROM editorial.concepts c
            WHERE c.importance_level IN ('forte', 'pilar')
        )
        SELECT * FROM concept_stats
        WHERE content_count >= 3
        ORDER BY (content_count - chapter_source_count) DESC, importance_score DESC
        LIMIT %s
    """, (batch_size,))
    return cursor.fetchall()


def infer_preferred_chapter_size(cursor, default=15):
    """Sinal de estilo (ADR-019 §5, escopo aprovado: só estrutura) — capítulos já
    revisados por um humano indicam o tamanho típico que costuma funcionar. Nunca
    influencia texto, só quantas fontes uma sugestão nova tenta reunir."""
    cursor.execute("""
        SELECT AVG(source_count)::int AS avg_size FROM (
            SELECT ch.id, COUNT(cs.id) AS source_count
            FROM editorial.chapters ch
            JOIN editorial.chapter_sources cs ON cs.chapter_id = ch.id
            WHERE ch.status = 'reviewed'
            GROUP BY ch.id
        ) sub
    """)
    row = cursor.fetchone()
    return row["avg_size"] if row and row["avg_size"] else default


def generate_chapter_suggestions_batch(cursor, batch_size=5):
    """Gera um lote de sugestões de capítulo (ADR-019 §4-5). Nunca publica nada — só
    grava em chapter_suggestions com status 'suggested', para revisão humana via
    /chapter-suggestions/<id>/promote."""
    candidates = find_chapter_suggestion_candidates(cursor, batch_size)
    preferred_size = infer_preferred_chapter_size(cursor)
    created = []

    for candidate in candidates:
        concept_id = str(candidate["id"])
        kept = build_proposed_sources(cursor, [concept_id], limit=preferred_size)
        if len(kept) < 2:
            continue

        source_titles = []
        proposed_sources = []
        for kind, item_id in kept:
            if kind == "segment":
                cursor.execute("SELECT title FROM editorial.content_segments WHERE id = %s", (item_id,))
                row = cursor.fetchone()
                source_titles.append(row["title"] if row else "(segmento)")
                proposed_sources.append({
                    "segment_id": str(item_id), "knowledge_card_id": None,
                    "inclusion_type": "literal_segment", "content": None,
                })
            else:
                cursor.execute("SELECT title FROM editorial.knowledge_cards WHERE id = %s", (item_id,))
                row = cursor.fetchone()
                source_titles.append(row["title"] if row else "(ficha)")
                proposed_sources.append({
                    "segment_id": None, "knowledge_card_id": str(item_id),
                    "inclusion_type": "card_synthesis", "content": None,
                })

        title, summary = generate_chapter_suggestion_label(candidate["canonical_name"], source_titles)
        if not title or not summary:
            continue

        cursor.execute("""
            INSERT INTO editorial.chapter_suggestions
                (title, summary, thematic_scope, proposed_sources, model)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING *
        """, (title, summary, Json([concept_id]), Json(proposed_sources), CHAT_MODEL))
        created.append(cursor.fetchone())

    return created


def require_api_key(function):
    @wraps(function)
    def decorated(*args, **kwargs):
        supplied_key = request.headers.get("X-Service-Api-Key", "")
        if not supplied_key or not hmac.compare_digest(supplied_key, SERVICE_API_KEY):
            return jsonify({"error": "Invalid service API key"}), 401
        return function(*args, **kwargs)
    return decorated


NORMALIZE_CONCEPT_KEY_SQL = "editorial.immutable_unaccent(lower(regexp_replace(trim(%s), '[-_]+', ' ', 'g')))"


def strip_embedding(row):
    """Remove o vetor de embedding de uma linha antes de serializar em JSON
    (tipo vector do pgvector não é serializável e o vetor não é útil ao cliente)."""
    if row is not None:
        row.pop("embedding", None)
    return row


def cosine_similarity(a, b):
    """Similaridade de cosseno entre dois embeddings (ADR-013 §3 — dedup de candidatos)."""
    if a is None or b is None:
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(y * y for y in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def json_safe(value):
    if isinstance(value, uuid.UUID):
        return str(value)
    if isinstance(value, datetime):
        return value.isoformat()
    return value


def serialize_rows(rows):
    return [{key: json_safe(value) for key, value in dict(row).items()} for row in rows]


def build_chapter_detail(cursor, chapter_id):
    """Monta o capítulo com o texto das fontes resolvido (ADR-013)."""
    cursor.execute("""
        SELECT ch.*, bp.title AS book_project_title
        FROM editorial.chapters ch
        JOIN editorial.book_projects bp ON bp.id = ch.book_project_id
        WHERE ch.id = %s
    """, (chapter_id,))
    chapter = cursor.fetchone()
    if not chapter:
        return None

    cursor.execute("""
        SELECT
            cs.id, cs.source_order, cs.inclusion_type, cs.content,
            cs.segment_id, cs.knowledge_card_id,
            seg.title AS segment_title, seg.full_text AS segment_full_text,
            seg.segment_type, seg.is_channeled,
            card.title AS card_title, card.summary AS card_summary
        FROM editorial.chapter_sources cs
        LEFT JOIN editorial.content_segments seg ON seg.id = cs.segment_id
        LEFT JOIN editorial.knowledge_cards card ON card.id = cs.knowledge_card_id
        WHERE cs.chapter_id = %s
        ORDER BY cs.source_order
    """, (chapter_id,))
    chapter["sources"] = cursor.fetchall()
    return chapter


def _fetch_chapter_sources(cursor, chapter_id):
    cursor.execute(
        "SELECT * FROM editorial.chapter_sources WHERE chapter_id = %s ORDER BY source_order",
        (chapter_id,),
    )
    return cursor.fetchall()


def snapshot_chapter_revision(cursor, chapter, existing_sources):
    """Grava a revisão do capítulo antes de suas fontes serem substituídas (ADR-013 §4)."""
    if not existing_sources:
        return
    cursor.execute("""
        INSERT INTO editorial.chapter_revisions (chapter_id, title, thematic_scope, status, sources_snapshot)
        VALUES (%(chapter_id)s, %(title)s, %(thematic_scope)s, %(status)s, %(sources_snapshot)s)
    """, {
        "chapter_id": chapter["id"],
        "title": chapter["title"],
        "thematic_scope": Json(chapter["thematic_scope"]),
        "status": chapter["status"],
        "sources_snapshot": Json(serialize_rows(existing_sources)),
    })
    cursor.execute("DELETE FROM editorial.chapter_sources WHERE chapter_id = %s", (chapter["id"],))


def resolve_concept_ids(cursor, concept_names):
    """Resolve nomes de conceito para IDs (ADR-011). Retorna (ids, nomes_desconhecidos) —
    nunca cria conceito novo aqui (diferente de sync_concepts, usado na destilação)."""
    concept_ids = []
    unknown = []
    for name in concept_names or []:
        cursor.execute(
            f"SELECT id FROM editorial.concepts WHERE normalized_key = {NORMALIZE_CONCEPT_KEY_SQL}",
            (name,),
        )
        row = cursor.fetchone()
        if row:
            concept_ids.append(str(row["id"]))
        else:
            unknown.append(name)
    return concept_ids, unknown


def insert_chapter(cursor, book_project_id, chapter_order, title, concept_ids):
    """Cria um capítulo (ADR-013). Extraído para ser reutilizável pela promoção de
    sugestões (ADR-019 §2), que precisa disso na mesma transação que insere as sources."""
    cursor.execute("""
        INSERT INTO editorial.chapters (book_project_id, chapter_order, title, thematic_scope)
        VALUES (%s, %s, %s, %s)
        RETURNING *
    """, (book_project_id, chapter_order, title, Json(concept_ids)))
    return cursor.fetchone()


def insert_chapter_sources(cursor, chapter_id, sources):
    """Insere chapter_sources a partir de uma lista de dicts (ADR-013) — mesmo formato
    usado por PUT /chapters/<id>/sources e pelos snapshots de chapter_revisions/
    chapter_suggestions.proposed_sources. Não apaga sources existentes nem gera
    revisão — isso é responsabilidade de quem chama (ver snapshot_chapter_revision)."""
    for order, item in enumerate(sources, start=1):
        cursor.execute("""
            INSERT INTO editorial.chapter_sources
                (chapter_id, segment_id, knowledge_card_id, source_order, inclusion_type, content)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (
            chapter_id, item.get("segment_id"), item.get("knowledge_card_id"),
            order, item["inclusion_type"], item.get("content"),
        ))


def sync_concepts(cursor, concept_texts, segment_id=None, card_id=None):
    """Resolve texto livre de conceitos para editorial.concepts (ADR-011) e grava os vínculos."""
    for raw_text in concept_texts or []:
        text = (raw_text or "").strip()
        if not text:
            continue
        cursor.execute("""
            INSERT INTO editorial.concepts (canonical_name)
            VALUES (%s)
            ON CONFLICT (normalized_key) DO UPDATE SET last_observed_at = NOW()
            RETURNING id
        """, (text,))
        concept_id = cursor.fetchone()["id"]
        if segment_id:
            cursor.execute("""
                INSERT INTO editorial.segment_concepts (segment_id, concept_id)
                VALUES (%s, %s)
                ON CONFLICT DO NOTHING
            """, (segment_id, concept_id))
        if card_id:
            cursor.execute("""
                INSERT INTO editorial.card_concepts (card_id, concept_id)
                VALUES (%s, %s)
                ON CONFLICT DO NOTHING
            """, (card_id, concept_id))


@app.get("/health")
def health():
    try:
        with get_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute("SELECT 1")
                cursor.fetchone()
        return jsonify({"status": "ok", "database": "connected"})
    except Exception as error:
        app.logger.exception("Health check failed")
        return jsonify({"status": "error", "database": "unavailable", "message": str(error)}), 500


@app.get("/themes")
@require_api_key
def list_themes():
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, theme_key, name, description, minimum_relevance, active
                FROM editorial.themes
                WHERE active = TRUE
                ORDER BY name
            """)
            rows = cursor.fetchall()
    return jsonify(rows)


@app.get("/segment-types")
@require_api_key
def list_segment_types():
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT type_key, name, description, editorially_disposable
                FROM editorial.segment_types
                ORDER BY name
            """)
            rows = cursor.fetchall()
    return jsonify(rows)


@app.get("/speakers")
@require_api_key
def list_speakers():
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, speaker_key, canonical_name, aliases, description, is_channeled_entity
                FROM editorial.speakers
                ORDER BY canonical_name
            """)
            rows = cursor.fetchall()
    return jsonify(rows)


@app.get("/concepts")
@require_api_key
def list_concepts():
    limit = min(request.args.get("limit", default=100, type=int), 500)
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT
                    c.id, c.canonical_name, c.aliases, c.description,
                    c.importance_score, c.importance_level, c.scope,
                    COUNT(DISTINCT sc.segment_id) AS segment_count,
                    COUNT(DISTINCT cc.card_id) AS card_count
                FROM editorial.concepts c
                LEFT JOIN editorial.segment_concepts sc ON sc.concept_id = c.id
                LEFT JOIN editorial.card_concepts cc ON cc.concept_id = c.id
                GROUP BY c.id
                ORDER BY c.importance_score DESC, c.canonical_name
                LIMIT %s
            """, (limit,))
            rows = cursor.fetchall()
    return jsonify(rows)


@app.post("/concepts/<concept_id>/scope")
@require_api_key
def set_concept_scope(concept_id):
    """Correção manual do escopo (ADR-022 §4) — a passada de enriquecimento marca
    conceitos como 'universal' automaticamente, mas o curador pode sempre reverter."""
    payload = request.get_json(silent=True) or {}
    scope = payload.get("scope")
    if scope not in ("universal", "tematico"):
        return jsonify({"error": "scope must be 'universal' or 'tematico'"}), 400
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                UPDATE editorial.concepts SET scope = %s
                WHERE id = %s
                RETURNING id, canonical_name, scope
            """, (scope, concept_id))
            row = cursor.fetchone()
    if not row:
        return jsonify({"error": "Concept not found"}), 404
    return jsonify(row)


@app.get("/concepts/<concept_id>/relations")
@require_api_key
def get_concept_relations(concept_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT id, canonical_name FROM editorial.concepts WHERE id = %s", (concept_id,))
            concept = cursor.fetchone()
            if not concept:
                return jsonify({"error": "Concept not found"}), 404

            cursor.execute("""
                SELECT
                    cr.id, cr.relation_type, cr.direction, cr.cooccurrence_count,
                    cr.first_observed_at, cr.last_observed_at,
                    other.id AS related_concept_id, other.canonical_name AS related_concept_name,
                    other.importance_score AS related_concept_importance_score
                FROM editorial.concept_relations cr
                JOIN editorial.concepts other
                  ON other.id = CASE WHEN cr.concept_a_id = %(id)s THEN cr.concept_b_id ELSE cr.concept_a_id END
                WHERE cr.concept_a_id = %(id)s OR cr.concept_b_id = %(id)s
                ORDER BY cr.cooccurrence_count DESC
            """, {"id": concept_id})
            relations = cursor.fetchall()

    return jsonify({"concept": concept, "relations": relations})


@app.post("/sources")
@require_api_key
def upsert_source():
    payload = request.get_json(silent=True) or {}
    required = ["external_file_id", "file_name", "source_type"]
    missing = [field for field in required if not payload.get(field)]
    if missing:
        return jsonify({"error": "Missing required fields", "fields": missing}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                INSERT INTO editorial.sources (
                    external_file_id, file_name, source_type, session_date,
                    drive_url, processing_status, updated_at
                ) VALUES (
                    %(external_file_id)s, %(file_name)s, %(source_type)s,
                    %(session_date)s, %(drive_url)s, %(processing_status)s, NOW()
                )
                ON CONFLICT (external_file_id)
                DO UPDATE SET
                    file_name = EXCLUDED.file_name,
                    source_type = EXCLUDED.source_type,
                    session_date = EXCLUDED.session_date,
                    drive_url = EXCLUDED.drive_url,
                    processing_status = EXCLUDED.processing_status,
                    updated_at = NOW()
                RETURNING *
            """, {
                "external_file_id": payload["external_file_id"],
                "file_name": payload["file_name"],
                "source_type": payload["source_type"],
                "session_date": payload.get("session_date"),
                "drive_url": payload.get("drive_url"),
                "processing_status": payload.get("processing_status", "pending"),
            })
            row = cursor.fetchone()
    return jsonify(row), 201


@app.post("/segments")
@require_api_key
def upsert_segment():
    payload = request.get_json(silent=True) or {}
    required = ["segment_key", "external_file_id", "segment_order", "segment_type", "title", "full_text"]
    missing = [field for field in required if payload.get(field) in (None, "")]
    if missing:
        return jsonify({"error": "Missing required fields", "fields": missing}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            # Precisa ser adquirido antes de qualquer INSERT/UPDATE desta transação, não só antes
            # de recalculate_concept_graph() — senão duas transações concorrentes podem tomar locks
            # de linha conflitantes (via sync_concepts/upsert) antes de disputar o advisory lock,
            # e o Postgres detecta um ciclo entre lock de linha e advisory lock (deadlock de novo).
            cursor.execute("SELECT pg_advisory_xact_lock(%s)", (CONCEPT_GRAPH_LOCK_KEY,))

            cursor.execute("SELECT id FROM editorial.sources WHERE external_file_id = %s", (payload["external_file_id"],))
            source = cursor.fetchone()
            if not source:
                return jsonify({"error": "Source not found", "external_file_id": payload["external_file_id"]}), 404

            cursor.execute("SELECT 1 FROM editorial.segment_types WHERE type_key = %s", (payload["segment_type"],))
            if not cursor.fetchone():
                return jsonify({"error": "Unknown segment_type", "segment_type": payload["segment_type"]}), 400

            cursor.execute("SELECT * FROM editorial.content_segments WHERE segment_key = %s", (payload["segment_key"],))
            existing = cursor.fetchone()
            if existing:
                # psycopg2 já desserializa colunas jsonb em listas/dicts Python ao buscar;
                # precisam ser reembrulhadas em Json() para voltar a ser gravadas como jsonb
                # (senão o psycopg2 as adapta como ARRAY do Postgres e a query falha).
                revision_params = dict(existing)
                for jsonb_field in ("keywords", "concepts", "related_themes", "editorial_applications"):
                    revision_params[jsonb_field] = Json(revision_params[jsonb_field])
                cursor.execute("""
                    INSERT INTO editorial.segment_revisions (
                        segment_id, segment_key, segment_order, segment_type, title,
                        executive_summary, full_text, keywords, concepts, related_themes,
                        editorial_applications, editorial_relevance, speaker_type, is_channeled
                    ) VALUES (
                        %(id)s, %(segment_key)s, %(segment_order)s, %(segment_type)s, %(title)s,
                        %(executive_summary)s, %(full_text)s, %(keywords)s, %(concepts)s, %(related_themes)s,
                        %(editorial_applications)s, %(editorial_relevance)s, %(speaker_type)s, %(is_channeled)s
                    )
                """, revision_params)

            speaker_key_lookup = payload.get("speaker_type") or "condutor_sessao"

            embedding_text = "\n\n".join(filter(None, [
                payload["title"], payload.get("executive_summary"), payload["full_text"],
            ]))
            embedding = generate_embedding(embedding_text)

            cursor.execute("""
                INSERT INTO editorial.content_segments (
                    segment_key, source_id, segment_order, segment_type, title,
                    executive_summary, full_text, keywords, concepts,
                    related_themes, editorial_applications, editorial_relevance,
                    speaker_type, is_channeled, speaker_id,
                    embedding, embedding_model, embedding_generated_at, updated_at
                ) VALUES (
                    %(segment_key)s, %(source_id)s, %(segment_order)s, %(segment_type)s,
                    %(title)s, %(executive_summary)s, %(full_text)s, %(keywords)s,
                    %(concepts)s, %(related_themes)s, %(editorial_applications)s,
                    %(editorial_relevance)s, %(speaker_type)s, %(is_channeled)s,
                    (SELECT id FROM editorial.speakers WHERE speaker_key = %(speaker_key_lookup)s),
                    %(embedding)s::vector, %(embedding_model)s, %(embedding_generated_at)s, NOW()
                )
                ON CONFLICT (segment_key)
                DO UPDATE SET
                    segment_order = EXCLUDED.segment_order,
                    segment_type = EXCLUDED.segment_type,
                    title = EXCLUDED.title,
                    executive_summary = EXCLUDED.executive_summary,
                    full_text = EXCLUDED.full_text,
                    keywords = EXCLUDED.keywords,
                    concepts = EXCLUDED.concepts,
                    related_themes = EXCLUDED.related_themes,
                    editorial_applications = EXCLUDED.editorial_applications,
                    editorial_relevance = EXCLUDED.editorial_relevance,
                    speaker_type = EXCLUDED.speaker_type,
                    is_channeled = EXCLUDED.is_channeled,
                    speaker_id = EXCLUDED.speaker_id,
                    embedding = EXCLUDED.embedding,
                    embedding_model = EXCLUDED.embedding_model,
                    embedding_generated_at = EXCLUDED.embedding_generated_at,
                    updated_at = NOW()
                RETURNING *
            """, {
                "segment_key": payload["segment_key"],
                "source_id": source["id"],
                "segment_order": payload["segment_order"],
                "segment_type": payload["segment_type"],
                "title": payload["title"],
                "executive_summary": payload.get("executive_summary"),
                "full_text": payload["full_text"],
                "keywords": Json(payload.get("keywords", [])),
                "concepts": Json(payload.get("concepts", [])),
                "related_themes": Json(payload.get("related_themes", [])),
                "editorial_applications": Json(payload.get("editorial_applications", [])),
                "editorial_relevance": payload.get("editorial_relevance", 0),
                "speaker_type": payload.get("speaker_type"),
                "is_channeled": bool(payload.get("is_channeled", False)),
                "speaker_key_lookup": speaker_key_lookup,
                "embedding": embedding,
                "embedding_model": EMBEDDING_MODEL if embedding else None,
                "embedding_generated_at": datetime.now(timezone.utc) if embedding else None,
            })
            row = cursor.fetchone()

            sync_concepts(cursor, payload.get("concepts", []), segment_id=row["id"])
            cursor.execute("SELECT editorial.recalculate_concept_graph()")

    return jsonify(strip_embedding(row)), 201


@app.get("/segments")
@require_api_key
def list_segments():
    theme = request.args.get("theme")
    segment_type = request.args.get("type")
    channeled = request.args.get("channeled")
    concept = request.args.get("concept")
    limit = min(request.args.get("limit", default=50, type=int), 500)

    clauses = []
    parameters = []
    if theme:
        clauses.append("cs.related_themes ? %s")
        parameters.append(theme)
    if segment_type:
        clauses.append("cs.segment_type = %s")
        parameters.append(segment_type)
    if channeled is not None:
        clauses.append("cs.is_channeled = %s")
        parameters.append(channeled.lower() == "true")
    if concept:
        clauses.append(f"""
            cs.id IN (
                SELECT sc.segment_id FROM editorial.segment_concepts sc
                JOIN editorial.concepts co ON co.id = sc.concept_id
                WHERE co.normalized_key = {NORMALIZE_CONCEPT_KEY_SQL}
            )
        """)
        parameters.append(concept)

    where = "WHERE " + " AND ".join(clauses) if clauses else ""
    parameters.append(limit)

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute(f"""
                SELECT cs.*, s.external_file_id, s.file_name AS source_file_name, s.session_date
                FROM editorial.content_segments cs
                JOIN editorial.sources s ON s.id = cs.source_id
                {where}
                ORDER BY cs.editorial_relevance DESC, s.session_date, cs.segment_order
                LIMIT %s
            """, parameters)
            rows = cursor.fetchall()
    return jsonify([strip_embedding(row) for row in rows])


@app.get("/segments/<segment_id>")
@require_api_key
def get_segment(segment_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT cs.*, s.external_file_id, s.file_name AS source_file_name, s.session_date
                FROM editorial.content_segments cs
                JOIN editorial.sources s ON s.id = cs.source_id
                WHERE cs.id = %s
            """, (segment_id,))
            row = cursor.fetchone()
    if not row:
        return jsonify({"error": "Segment not found"}), 404
    return jsonify(strip_embedding(row))


@app.get("/segment-insights")
@require_api_key
def list_all_segment_insights():
    """Lista cartões de insight de todo o acervo, não escopados a um segmento (ADR-023,
    tela Explorar Acervo) -- só os que já foram gerados, nunca dispara geração nova."""
    status_filter = request.args.get("status")
    limit = min(request.args.get("limit", default=50, type=int), 500)
    clauses = []
    parameters = []
    if status_filter:
        clauses.append("si.status = %s")
        parameters.append(status_filter)
    else:
        clauses.append("si.status != 'dismissed'")
    where = "WHERE " + " AND ".join(clauses)
    parameters.append(limit)

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute(f"""
                SELECT si.*, seg.title AS segment_title, seg.is_channeled
                FROM editorial.segment_insights si
                JOIN editorial.content_segments seg ON seg.id = si.segment_id
                {where}
                ORDER BY si.generated_at DESC
                LIMIT %s
            """, parameters)
            rows = cursor.fetchall()
    return jsonify(rows)


@app.get("/segments/<segment_id>/insights")
@require_api_key
def list_segment_insights(segment_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT id FROM editorial.content_segments WHERE id = %s", (segment_id,))
            if not cursor.fetchone():
                return jsonify({"error": "Segment not found"}), 404
            cursor.execute("""
                SELECT * FROM editorial.segment_insights
                WHERE segment_id = %s
                ORDER BY generated_at
            """, (segment_id,))
            rows = cursor.fetchall()
    return jsonify(rows)


def generate_and_persist_segment_insights(cursor, segment_id, segment_full_text):
    """Núcleo compartilhado da geração de cartão de insight (ADR-022 §2) — usado tanto
    pela geração sob demanda (um segmento) quanto pelo lote controlado (ADR-023 §3)."""
    # Mesma disputa de locks que upsert_segment/upsert_knowledge_card (ver comentário lá):
    # esta função também escreve em editorial.concepts/concept_relations, que
    # recalculate_concept_graph() varre por inteiro.
    cursor.execute("SELECT pg_advisory_xact_lock(%s)", (CONCEPT_GRAPH_LOCK_KEY,))

    cursor.execute("""
        SELECT c.id, c.canonical_name
        FROM editorial.concepts c
        JOIN editorial.segment_concepts sc ON sc.concept_id = c.id
        WHERE sc.segment_id = %s
    """, (segment_id,))
    existing_names = [row["canonical_name"] for row in cursor.fetchall()]

    insights, relations = generate_segment_insights(segment_full_text, existing_names)
    if not insights:
        return []

    all_related_names = set()
    for insight in insights:
        all_related_names.update(insight.get("related_concepts") or [])
    sync_concepts(cursor, sorted(all_related_names), segment_id=segment_id)

    cursor.execute("""
        SELECT c.id, c.canonical_name
        FROM editorial.concepts c
        JOIN editorial.segment_concepts sc ON sc.concept_id = c.id
        WHERE sc.segment_id = %s
    """, (segment_id,))
    name_to_id = {row["canonical_name"]: row["id"] for row in cursor.fetchall()}

    created = []
    for insight in insights:
        cursor.execute("""
            INSERT INTO editorial.segment_insights (
                segment_id, concept_title, explanation, philosophical_context,
                practical_application, related_concepts, model
            ) VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING *
        """, (
            segment_id,
            (insight.get("concept_title") or "")[:500],
            insight.get("explanation") or "",
            insight.get("philosophical_context") or "",
            insight.get("practical_application") or "",
            Json(insight.get("related_concepts") or []),
            CHAT_MODEL,
        ))
        created.append(cursor.fetchone())

    if name_to_id:
        cursor.execute("""
            UPDATE editorial.concepts
            SET scope = 'universal'
            WHERE id = ANY(%s::uuid[]) AND scope = 'tematico'
        """, (list(name_to_id.values()),))

    apply_relation_types(cursor, name_to_id, relations)
    return created


@app.post("/segments/<segment_id>/insights")
@require_api_key
def generate_segment_insight_cards(segment_id):
    """Gera cartões de insight sob demanda (ADR-022 §2) — nunca em lote automático.
    Nasce sempre em status 'suggested': o curador decide revisar ou descartar."""
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT id, full_text FROM editorial.content_segments WHERE id = %s", (segment_id,))
            segment = cursor.fetchone()
            if not segment:
                return jsonify({"error": "Segment not found"}), 404

            created = generate_and_persist_segment_insights(cursor, segment_id, segment["full_text"])
            if not created:
                return jsonify({"error": "Failed to generate insights"}), 502

    return jsonify(serialize_rows(created)), 201


@app.post("/segment-insights/generate-batch")
@require_api_key
def generate_segment_insights_batch():
    """Lote explícito e controlado (ADR-023 §3) — nunca automático/agendado. Processa os
    N segmentos canalizados de maior editorial_relevance que ainda não têm nenhum
    segment_insight. Mantém a cautela de custo da ADR-022 §2: só roda quando alguém pede."""
    payload = request.get_json(silent=True) or {}
    batch_size = min(int(payload.get("batch_size", 5)), 20)

    results = {"processed": [], "errors": []}
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT cs.id, cs.title, cs.full_text
                FROM editorial.content_segments cs
                WHERE cs.is_channeled = TRUE
                  AND NOT EXISTS (
                      SELECT 1 FROM editorial.segment_insights si WHERE si.segment_id = cs.id
                  )
                ORDER BY cs.editorial_relevance DESC
                LIMIT %s
            """, (batch_size,))
            candidates = cursor.fetchall()

            for segment in candidates:
                try:
                    created = generate_and_persist_segment_insights(cursor, segment["id"], segment["full_text"])
                    connection.commit()
                    results["processed"].append({
                        "segment_id": segment["id"],
                        "segment_title": segment["title"],
                        "insights_created": len(created),
                    })
                except Exception as error:
                    connection.rollback()
                    results["errors"].append({"segment_id": segment["id"], "reason": str(error)})

    return jsonify(results)


@app.post("/segment-insights/<insight_id>/review")
@require_api_key
def review_segment_insight(insight_id):
    """Gate humano (ADR-022 §2) — marca um cartão de insight como revisado pelo curador."""
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                UPDATE editorial.segment_insights SET status = 'reviewed'
                WHERE id = %s
                RETURNING *
            """, (insight_id,))
            row = cursor.fetchone()
    if not row:
        return jsonify({"error": "Segment insight not found"}), 404
    return jsonify(row)


@app.post("/segment-insights/<insight_id>/dismiss")
@require_api_key
def dismiss_segment_insight(insight_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                UPDATE editorial.segment_insights SET status = 'dismissed'
                WHERE id = %s
                RETURNING *
            """, (insight_id,))
            row = cursor.fetchone()
    if not row:
        return jsonify({"error": "Segment insight not found"}), 404
    return jsonify(row)


@app.post("/knowledge-cards")
@require_api_key
def upsert_knowledge_card():
    payload = request.get_json(silent=True) or {}
    required = ["card_key", "external_file_id", "theme_key", "title", "summary", "relevance_score"]
    missing = [field for field in required if payload.get(field) is None]
    if missing:
        return jsonify({"error": "Missing required fields", "fields": missing}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            # Ver comentário equivalente em upsert_segment: precisa ser o primeiro comando da
            # transação, antes de qualquer INSERT/UPDATE.
            cursor.execute("SELECT pg_advisory_xact_lock(%s)", (CONCEPT_GRAPH_LOCK_KEY,))

            cursor.execute("SELECT id FROM editorial.sources WHERE external_file_id = %s", (payload["external_file_id"],))
            source = cursor.fetchone()
            if not source:
                return jsonify({"error": "Source not found", "external_file_id": payload["external_file_id"]}), 404

            cursor.execute("SELECT id FROM editorial.themes WHERE theme_key = %s AND active = TRUE", (payload["theme_key"],))
            theme = cursor.fetchone()
            if not theme:
                return jsonify({"error": "Active theme not found", "theme_key": payload["theme_key"]}), 404

            segment_id = None
            if payload.get("segment_key"):
                cursor.execute("SELECT id FROM editorial.content_segments WHERE segment_key = %s", (payload["segment_key"],))
                segment = cursor.fetchone()
                if segment:
                    segment_id = segment["id"]

            cursor.execute("SELECT * FROM editorial.knowledge_cards WHERE card_key = %s", (payload["card_key"],))
            existing = cursor.fetchone()
            if existing:
                # ver comentário equivalente em upsert_segment: jsonb já vem desserializado
                # do psycopg2, precisa ser reembrulhado em Json() antes de regravar.
                revision_params = dict(existing)
                for jsonb_field in ("concepts", "principles", "quotes", "evidence"):
                    revision_params[jsonb_field] = Json(revision_params[jsonb_field])
                cursor.execute("""
                    INSERT INTO editorial.card_revisions (
                        card_id, card_key, title, summary, concepts, principles,
                        quotes, evidence, relevance_score
                    ) VALUES (
                        %(id)s, %(card_key)s, %(title)s, %(summary)s, %(concepts)s, %(principles)s,
                        %(quotes)s, %(evidence)s, %(relevance_score)s
                    )
                """, revision_params)

            embedding_text = "\n\n".join(filter(None, [payload["title"], payload["summary"]]))
            embedding = generate_embedding(embedding_text)

            cursor.execute("""
                INSERT INTO editorial.knowledge_cards (
                    card_key, source_id, theme_id, segment_id, block_number,
                    title, summary, concepts, principles, quotes, evidence,
                    relevance_score, embedding, embedding_model, embedding_generated_at, updated_at
                ) VALUES (
                    %(card_key)s, %(source_id)s, %(theme_id)s, %(segment_id)s,
                    %(block_number)s, %(title)s, %(summary)s, %(concepts)s,
                    %(principles)s, %(quotes)s, %(evidence)s,
                    %(relevance_score)s, %(embedding)s::vector, %(embedding_model)s, %(embedding_generated_at)s, NOW()
                )
                ON CONFLICT (card_key)
                DO UPDATE SET
                    segment_id = EXCLUDED.segment_id,
                    title = EXCLUDED.title,
                    summary = EXCLUDED.summary,
                    concepts = EXCLUDED.concepts,
                    principles = EXCLUDED.principles,
                    quotes = EXCLUDED.quotes,
                    evidence = EXCLUDED.evidence,
                    relevance_score = EXCLUDED.relevance_score,
                    embedding = EXCLUDED.embedding,
                    embedding_model = EXCLUDED.embedding_model,
                    embedding_generated_at = EXCLUDED.embedding_generated_at,
                    updated_at = NOW()
                RETURNING *
            """, {
                "card_key": payload["card_key"],
                "source_id": source["id"],
                "theme_id": theme["id"],
                "segment_id": segment_id,
                "block_number": payload.get("block_number"),
                "title": payload["title"],
                "summary": payload["summary"],
                "concepts": Json(payload.get("concepts", [])),
                "principles": Json(payload.get("principles", [])),
                "quotes": Json(payload.get("quotes", [])),
                "evidence": Json(payload.get("evidence", [])),
                "relevance_score": payload["relevance_score"],
                "embedding": embedding,
                "embedding_model": EMBEDDING_MODEL if embedding else None,
                "embedding_generated_at": datetime.now(timezone.utc) if embedding else None,
            })
            row = cursor.fetchone()

            sync_concepts(cursor, payload.get("concepts", []), card_id=row["id"])
            cursor.execute("SELECT editorial.recalculate_concept_graph()")

    return jsonify(strip_embedding(row)), 201


@app.get("/knowledge-cards")
@require_api_key
def list_knowledge_cards():
    theme_key = request.args.get("theme")
    concept = request.args.get("concept")
    limit = min(request.args.get("limit", default=50, type=int), 500)
    clauses = []
    parameters = []
    if theme_key:
        clauses.append("t.theme_key = %s")
        parameters.append(theme_key)
    if concept:
        clauses.append(f"""
            kc.id IN (
                SELECT cc.card_id FROM editorial.card_concepts cc
                JOIN editorial.concepts co ON co.id = cc.concept_id
                WHERE co.normalized_key = {NORMALIZE_CONCEPT_KEY_SQL}
            )
        """)
        parameters.append(concept)
    where = "WHERE " + " AND ".join(clauses) if clauses else ""
    parameters.append(limit)

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute(f"""
                SELECT
                    kc.id, kc.card_key, kc.title, kc.summary, kc.concepts,
                    kc.principles, kc.quotes, kc.evidence, kc.relevance_score,
                    kc.importance_score, kc.importance_level, kc.segment_id,
                    s.external_file_id, s.file_name AS source_file_name,
                    s.session_date, t.theme_key, t.name AS theme_name
                FROM editorial.knowledge_cards kc
                JOIN editorial.sources s ON s.id = kc.source_id
                JOIN editorial.themes t ON t.id = kc.theme_id
                {where}
                ORDER BY kc.importance_score DESC, kc.created_at DESC
                LIMIT %s
            """, parameters)
            rows = cursor.fetchall()
    return jsonify(rows)


@app.post("/reindex-embeddings")
@require_api_key
def reindex_embeddings():
    """Gera embeddings para segmentos/fichas que ainda não têm (backfill de dados
    anteriores à ADR-012, ou reprocessamento se o modelo de embedding mudar no futuro).
    Passe {"force": true} para regerar mesmo os que já têm embedding."""
    payload = request.get_json(silent=True) or {}
    force = bool(payload.get("force", False))
    limit = min(int(payload.get("limit", 200)), 500)

    updated = {"segments": 0, "knowledge_cards": 0}
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            where = "" if force else "WHERE embedding IS NULL"
            cursor.execute(f"""
                SELECT id, title, executive_summary, full_text
                FROM editorial.content_segments
                {where}
                ORDER BY created_at
                LIMIT %s
            """, (limit,))
            segments = cursor.fetchall()

            for segment in segments:
                text = "\n\n".join(filter(None, [
                    segment["title"], segment["executive_summary"], segment["full_text"],
                ]))
                embedding = generate_embedding(text)
                if embedding is None:
                    continue
                cursor.execute("""
                    UPDATE editorial.content_segments
                    SET embedding = %s::vector, embedding_model = %s, embedding_generated_at = NOW()
                    WHERE id = %s
                """, (embedding, EMBEDDING_MODEL, segment["id"]))
                updated["segments"] += 1

            cursor.execute(f"""
                SELECT id, title, summary
                FROM editorial.knowledge_cards
                {where}
                ORDER BY created_at
                LIMIT %s
            """, (limit,))
            cards = cursor.fetchall()

            for card in cards:
                text = "\n\n".join(filter(None, [card["title"], card["summary"]]))
                embedding = generate_embedding(text)
                if embedding is None:
                    continue
                cursor.execute("""
                    UPDATE editorial.knowledge_cards
                    SET embedding = %s::vector, embedding_model = %s, embedding_generated_at = NOW()
                    WHERE id = %s
                """, (embedding, EMBEDDING_MODEL, card["id"]))
                updated["knowledge_cards"] += 1

    return jsonify(updated)


@app.post("/search")
@require_api_key
def semantic_search():
    """Busca semântica (ADR-012 §5): embeda a consulta e ranqueia segmentos/fichas por
    similaridade de cosseno. Não substitui os filtros exatos de /segments e /knowledge-cards,
    complementa quando a mesma ideia aparece com vocabulário diferente ao longo do acervo."""
    payload = request.get_json(silent=True) or {}
    query = (payload.get("query") or "").strip()
    if not query:
        return jsonify({"error": "Missing required field", "fields": ["query"]}), 400

    limit = min(int(payload.get("limit", 10)), 50)
    include = payload.get("include") or ["segments", "knowledge_cards"]

    query_embedding = generate_embedding(query)
    if query_embedding is None:
        return jsonify({"error": "Failed to generate embedding for query"}), 502

    results = []
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            if "segments" in include:
                cursor.execute("""
                    SELECT
                        'segment' AS result_type, cs.id, cs.title, cs.executive_summary,
                        cs.segment_type, cs.editorial_relevance, cs.is_channeled,
                        s.external_file_id, s.file_name AS source_file_name, s.session_date,
                        1 - (cs.embedding <=> %(embedding)s::vector) AS similarity
                    FROM editorial.content_segments cs
                    JOIN editorial.sources s ON s.id = cs.source_id
                    WHERE cs.embedding IS NOT NULL
                    ORDER BY cs.embedding <=> %(embedding)s::vector
                    LIMIT %(limit)s
                """, {"embedding": query_embedding, "limit": limit})
                results.extend(cursor.fetchall())

            if "knowledge_cards" in include:
                cursor.execute("""
                    SELECT
                        'knowledge_card' AS result_type, kc.id, kc.title, kc.summary,
                        kc.importance_score, kc.importance_level,
                        s.external_file_id, s.file_name AS source_file_name, s.session_date,
                        1 - (kc.embedding <=> %(embedding)s::vector) AS similarity
                    FROM editorial.knowledge_cards kc
                    JOIN editorial.sources s ON s.id = kc.source_id
                    WHERE kc.embedding IS NOT NULL
                    ORDER BY kc.embedding <=> %(embedding)s::vector
                    LIMIT %(limit)s
                """, {"embedding": query_embedding, "limit": limit})
                results.extend(cursor.fetchall())

    results.sort(key=lambda item: item["similarity"], reverse=True)
    return jsonify(results[:limit])


@app.post("/book-projects")
@require_api_key
def create_book_project():
    payload = request.get_json(silent=True) or {}
    if not payload.get("title"):
        return jsonify({"error": "Missing required fields", "fields": ["title"]}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                INSERT INTO editorial.book_projects (title, description)
                VALUES (%s, %s)
                RETURNING *
            """, (payload["title"], payload.get("description")))
            row = cursor.fetchone()
    return jsonify(row), 201


@app.get("/book-projects")
@require_api_key
def list_book_projects():
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT bp.*, count(ch.id) AS chapter_count
                FROM editorial.book_projects bp
                LEFT JOIN editorial.chapters ch ON ch.book_project_id = bp.id
                GROUP BY bp.id
                ORDER BY bp.created_at DESC
            """)
            rows = cursor.fetchall()
    return jsonify(rows)


@app.get("/book-projects/<book_project_id>")
@require_api_key
def get_book_project(book_project_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.book_projects WHERE id = %s", (book_project_id,))
            project = cursor.fetchone()
            if not project:
                return jsonify({"error": "Book project not found"}), 404

            cursor.execute("""
                SELECT ch.id, ch.chapter_order, ch.title, ch.status, ch.thematic_scope,
                       count(cs.id) AS source_count
                FROM editorial.chapters ch
                LEFT JOIN editorial.chapter_sources cs ON cs.chapter_id = ch.id
                WHERE ch.book_project_id = %s
                GROUP BY ch.id
                ORDER BY ch.chapter_order
            """, (book_project_id,))
            project["chapters"] = cursor.fetchall()
    return jsonify(project)


@app.post("/book-projects/<book_project_id>/chapters")
@require_api_key
def create_chapter(book_project_id):
    payload = request.get_json(silent=True) or {}
    required = ["title", "chapter_order"]
    missing = [field for field in required if payload.get(field) in (None, "")]
    if missing:
        return jsonify({"error": "Missing required fields", "fields": missing}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT id FROM editorial.book_projects WHERE id = %s", (book_project_id,))
            if not cursor.fetchone():
                return jsonify({"error": "Book project not found"}), 404

            concept_ids, unknown_concepts = resolve_concept_ids(cursor, payload.get("thematic_scope", []))
            if unknown_concepts:
                return jsonify({"error": "Unknown concepts", "concepts": unknown_concepts}), 400

            chapter = insert_chapter(cursor, book_project_id, payload["chapter_order"], payload["title"], concept_ids)
    return jsonify(chapter), 201


@app.get("/chapters/<chapter_id>")
@require_api_key
def get_chapter(chapter_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            chapter = build_chapter_detail(cursor, chapter_id)
    if not chapter:
        return jsonify({"error": "Chapter not found"}), 404
    return jsonify(chapter)


@app.post("/chapters/<chapter_id>/propose")
@require_api_key
def propose_chapter_sources(chapter_id):
    """Módulo 06 (ADR-013 §3): gera uma PROPOSTA de montagem, não uma versão final.
    Recupera candidatos por conceito (Mapa Filosófico), ranqueia e remove redundância
    por similaridade de embedding. Sempre grava como proposta — o capítulo permanece em
    'draft' até aprovação humana explícita via POST /chapters/<id>/approve."""
    payload = request.get_json(silent=True) or {}
    min_importance = int(payload.get("min_importance", 25))
    limit = min(int(payload.get("limit", 20)), 50)
    similarity_threshold = float(payload.get("similarity_threshold", 0.92))

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapters WHERE id = %s", (chapter_id,))
            chapter = cursor.fetchone()
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404

            concept_ids = chapter["thematic_scope"] or []
            if not concept_ids:
                return jsonify({"error": "Chapter has empty thematic_scope"}), 400

            kept = build_proposed_sources(cursor, concept_ids, limit, min_importance, similarity_threshold)

            snapshot_chapter_revision(
                cursor, chapter,
                _fetch_chapter_sources(cursor, chapter_id),
            )

            for order, (kind, item_id) in enumerate(kept, start=1):
                if kind == "segment":
                    cursor.execute("""
                        INSERT INTO editorial.chapter_sources (chapter_id, segment_id, source_order, inclusion_type)
                        VALUES (%s, %s, %s, 'literal_segment')
                    """, (chapter_id, item_id, order))
                else:
                    cursor.execute("""
                        INSERT INTO editorial.chapter_sources
                            (chapter_id, knowledge_card_id, source_order, inclusion_type)
                        VALUES (%s, %s, %s, 'card_synthesis')
                    """, (chapter_id, item_id, order))

            cursor.execute("UPDATE editorial.chapters SET updated_at = NOW() WHERE id = %s", (chapter_id,))
            result = build_chapter_detail(cursor, chapter_id)
    return jsonify(result)


@app.post("/chapter-suggestions/generate")
@require_api_key
def generate_chapter_suggestions():
    """Service-only (ADR-019 §6) — disparado pelo fluxo agendado do n8n, nunca pela
    UI. Nunca publica nada sozinho: só cria linhas em chapter_suggestions com status
    'suggested', para revisão humana via GET/POST /chapter-suggestions/*."""
    payload = request.get_json(silent=True) or {}
    batch_size = min(int(payload.get("batch_size", 5)), 20)
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            created = generate_chapter_suggestions_batch(cursor, batch_size)
    return jsonify(serialize_rows(created)), 201


@app.get("/chapter-suggestions")
@require_api_key
def list_chapter_suggestions():
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                SELECT id, title, summary, thematic_scope, status, generated_at
                FROM editorial.chapter_suggestions
                ORDER BY generated_at DESC
            """)
            rows = cursor.fetchall()
    return jsonify(rows)


@app.get("/chapter-suggestions/<suggestion_id>")
@require_api_key
def get_chapter_suggestion(suggestion_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapter_suggestions WHERE id = %s", (suggestion_id,))
            suggestion = cursor.fetchone()
            if not suggestion:
                return jsonify({"error": "Chapter suggestion not found"}), 404

            sources = []
            for item in suggestion["proposed_sources"]:
                if item.get("segment_id"):
                    cursor.execute("""
                        SELECT id, title AS segment_title, full_text AS segment_full_text,
                               segment_type, is_channeled
                        FROM editorial.content_segments WHERE id = %s
                    """, (item["segment_id"],))
                    sources.append({**item, **(cursor.fetchone() or {})})
                elif item.get("knowledge_card_id"):
                    cursor.execute("""
                        SELECT id, title AS card_title, summary AS card_summary
                        FROM editorial.knowledge_cards WHERE id = %s
                    """, (item["knowledge_card_id"],))
                    sources.append({**item, **(cursor.fetchone() or {})})
            suggestion["sources"] = sources
    return jsonify(suggestion)


@app.post("/chapter-suggestions/<suggestion_id>/promote")
@require_api_key
def promote_chapter_suggestion(suggestion_id):
    """Transação única (ADR-019 §3): cria o capítulo real, insere as sources do
    snapshot e marca a sugestão como promovida — tudo dentro da mesma conexão, nunca
    como chamadas HTTP encadeadas que poderiam ficar pela metade em caso de falha."""
    payload = request.get_json(silent=True) or {}
    book_project_id = payload.get("book_project_id")
    if not book_project_id:
        return jsonify({"error": "Missing required field", "fields": ["book_project_id"]}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapter_suggestions WHERE id = %s", (suggestion_id,))
            suggestion = cursor.fetchone()
            if not suggestion:
                return jsonify({"error": "Chapter suggestion not found"}), 404
            if suggestion["status"] != "suggested":
                return jsonify({"error": "Suggestion is not in suggested status", "status": suggestion["status"]}), 400

            cursor.execute("SELECT id FROM editorial.book_projects WHERE id = %s", (book_project_id,))
            if not cursor.fetchone():
                return jsonify({"error": "Book project not found"}), 404

            chapter_order = payload.get("chapter_order")
            if chapter_order is None:
                cursor.execute(
                    "SELECT COALESCE(MAX(chapter_order), 0) + 1 AS next_order FROM editorial.chapters WHERE book_project_id = %s",
                    (book_project_id,),
                )
                chapter_order = cursor.fetchone()["next_order"]

            title = payload.get("title") or suggestion["title"]
            concept_ids = suggestion["thematic_scope"] or []

            chapter = insert_chapter(cursor, book_project_id, chapter_order, title, concept_ids)
            insert_chapter_sources(cursor, chapter["id"], suggestion["proposed_sources"])

            cursor.execute("""
                UPDATE editorial.chapter_suggestions
                SET status = 'promoted', promoted_chapter_id = %s
                WHERE id = %s
            """, (chapter["id"], suggestion_id))

            result = build_chapter_detail(cursor, chapter["id"])
    return jsonify(result), 201


@app.post("/chapter-suggestions/<suggestion_id>/dismiss")
@require_api_key
def dismiss_chapter_suggestion(suggestion_id):
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("""
                UPDATE editorial.chapter_suggestions SET status = 'dismissed'
                WHERE id = %s AND status = 'suggested'
                RETURNING *
            """, (suggestion_id,))
            row = cursor.fetchone()
    if not row:
        return jsonify({"error": "Chapter suggestion not found or not in suggested status"}), 404
    return jsonify(row)


@app.put("/chapters/<chapter_id>/sources")
@require_api_key
def set_chapter_sources(chapter_id):
    """Substituição manual das fontes de um capítulo — para quando um humano quer
    ajustar/curar a proposta gerada por /propose antes de aprovar."""
    payload = request.get_json(silent=True) or {}
    sources = payload.get("sources")
    if not isinstance(sources, list) or not sources:
        return jsonify({"error": "Missing required field", "fields": ["sources"]}), 400

    for item in sources:
        inclusion_type = item.get("inclusion_type")
        if inclusion_type not in ("literal_segment", "card_synthesis", "transition_context"):
            return jsonify({"error": "Invalid inclusion_type", "value": inclusion_type}), 400
        if inclusion_type == "literal_segment" and not item.get("segment_id"):
            return jsonify({"error": "literal_segment requires segment_id"}), 400
        if inclusion_type == "card_synthesis" and not item.get("knowledge_card_id"):
            return jsonify({"error": "card_synthesis requires knowledge_card_id"}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapters WHERE id = %s", (chapter_id,))
            chapter = cursor.fetchone()
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404

            snapshot_chapter_revision(
                cursor, chapter,
                _fetch_chapter_sources(cursor, chapter_id),
            )

            insert_chapter_sources(cursor, chapter_id, sources)

            cursor.execute("UPDATE editorial.chapters SET updated_at = NOW() WHERE id = %s", (chapter_id,))
            result = build_chapter_detail(cursor, chapter_id)
    return jsonify(result)


def extract_literal_segments(doc):
    """Percorre um documento Tiptap (JSON) e retorna {segment_id: text} de todo nó
    literalSegment encontrado, em qualquer profundidade (ADR-020 Fase B)."""
    found = {}

    def walk(node):
        if not isinstance(node, dict):
            return
        if node.get("type") == "literalSegment":
            attrs = node.get("attrs") or {}
            segment_id = attrs.get("segmentId")
            if segment_id:
                found[segment_id] = attrs.get("text", "")
        for child in node.get("content") or []:
            walk(child)

    walk(doc)
    return found


def compose_manuscript_from_sources(cursor, chapter_id):
    """Monta o documento Tiptap inicial do manuscrito a partir de chapter_sources (usado só
    na primeira leitura, antes de qualquer manuscript_content existir). literal_segment vira
    nó travado; card_synthesis/transition_context viram parágrafos editáveis comuns."""
    chapter = build_chapter_detail(cursor, chapter_id)
    content = []
    for source in chapter["sources"]:
        if source["inclusion_type"] == "literal_segment":
            content.append({
                "type": "literalSegment",
                "attrs": {
                    "segmentId": str(source["segment_id"]),
                    "title": source["segment_title"],
                    "text": source["segment_full_text"],
                },
            })
        else:
            text = source["content"] if source["inclusion_type"] == "transition_context" else source["card_summary"]
            for paragraph in (text or "").split("\n\n"):
                if paragraph.strip():
                    content.append({"type": "paragraph", "content": [{"type": "text", "text": paragraph.strip()}]})
    return {"type": "doc", "content": content}


@app.get("/chapters/<chapter_id>/manuscript")
@require_api_key
def get_chapter_manuscript(chapter_id):
    """ADR-020 Fase B. Se o manuscrito ainda não foi salvo nenhuma vez, compõe uma versão
    inicial a partir das chapter_sources atuais (sem persistir — só a primeira gravação real
    via PUT formaliza o manuscrito)."""
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapters WHERE id = %s", (chapter_id,))
            chapter = cursor.fetchone()
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404

            if chapter["manuscript_content"] is not None:
                return jsonify({
                    "manuscript_content": chapter["manuscript_content"],
                    "manuscript_updated_at": chapter["manuscript_updated_at"],
                    "initialized": True,
                })

            composed = compose_manuscript_from_sources(cursor, chapter_id)
    return jsonify({"manuscript_content": composed, "manuscript_updated_at": None, "initialized": False})


@app.put("/chapters/<chapter_id>/manuscript")
@require_api_key
def set_chapter_manuscript(chapter_id):
    """ADR-020 Fase B. Nunca confia no cliente para preservar blocos literais: a cada
    gravação, verifica que todo segment_id exigido continua presente e que seu texto é
    idêntico a content_segments.full_text (fonte de verdade). Rejeita a gravação inteira se
    qualquer bloco travado tiver sido removido ou alterado."""
    payload = request.get_json(silent=True) or {}
    new_content = payload.get("manuscript_content")
    if not isinstance(new_content, dict):
        return jsonify({"error": "Missing required field", "fields": ["manuscript_content"]}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapters WHERE id = %s", (chapter_id,))
            chapter = cursor.fetchone()
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404

            new_segments = extract_literal_segments(new_content)

            if chapter["manuscript_content"] is not None:
                required_ids = set(extract_literal_segments(chapter["manuscript_content"]).keys())
            else:
                cursor.execute("""
                    SELECT segment_id::text AS segment_id FROM editorial.chapter_sources
                    WHERE chapter_id = %s AND inclusion_type = 'literal_segment' AND segment_id IS NOT NULL
                """, (chapter_id,))
                required_ids = {row["segment_id"] for row in cursor.fetchall()}

            missing = required_ids - set(new_segments.keys())
            if missing:
                return jsonify({
                    "error": "Manuscript is missing required literal segments",
                    "segment_ids": sorted(missing),
                }), 400

            if new_segments:
                cursor.execute(
                    "SELECT id::text AS id, full_text FROM editorial.content_segments WHERE id = ANY(%s::uuid[])",
                    (list(new_segments.keys()),),
                )
                truth = {row["id"]: row["full_text"] for row in cursor.fetchall()}
                for segment_id, text in new_segments.items():
                    if segment_id not in truth:
                        return jsonify({"error": "Unknown segment_id in manuscript", "segment_id": segment_id}), 400
                    if text != truth[segment_id]:
                        return jsonify({
                            "error": "Literal segment text does not match source of truth",
                            "segment_id": segment_id,
                        }), 400

            cursor.execute("""
                UPDATE editorial.chapters
                SET manuscript_content = %s, manuscript_updated_at = NOW()
                WHERE id = %s
            """, (Json(new_content), chapter_id))
    return jsonify({"ok": True})


@app.post("/chapters/<chapter_id>/manuscript/checkpoint")
@require_api_key
def checkpoint_chapter_manuscript(chapter_id):
    """Checkpoint explícito ("salvar versão") — nunca automático/silencioso (ADR-020 Fase B)."""
    payload = request.get_json(silent=True) or {}
    label = payload.get("label")

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT manuscript_content FROM editorial.chapters WHERE id = %s", (chapter_id,))
            chapter = cursor.fetchone()
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404
            if chapter["manuscript_content"] is None:
                return jsonify({"error": "Manuscript has no content to checkpoint yet"}), 400

            cursor.execute("""
                INSERT INTO editorial.chapter_manuscript_revisions (chapter_id, manuscript_content, label)
                VALUES (%s, %s, %s)
                RETURNING id, created_at
            """, (chapter_id, Json(chapter["manuscript_content"]), label))
            row = cursor.fetchone()
    return jsonify(row), 201


def build_chapter_docx(project_title, chapter_title, manuscript_content):
    """ADR-015 §4. Percorre o manuscrito (Tiptap JSON) e monta o .docx — blocos literalSegment
    viram parágrafos recuados/itálicos com o título da canalização como legenda; parágrafos
    comuns preservam negrito/itálico. Nunca lê chapter_sources diretamente."""
    document = Document()
    document.add_heading(chapter_title, level=1)
    subtitle = document.add_paragraph(project_title)
    subtitle.runs[0].italic = True

    for node in manuscript_content.get("content") or []:
        node_type = node.get("type")
        if node_type == "literalSegment":
            attrs = node.get("attrs") or {}
            caption = document.add_paragraph()
            caption_run = caption.add_run(f"Canalização — {attrs.get('title')}" if attrs.get("title") else "Canalização")
            caption_run.bold = True
            caption_run.italic = True
            caption_run.font.size = Pt(9)
            body = document.add_paragraph()
            body.paragraph_format.left_indent = Inches(0.3)
            body_run = body.add_run(attrs.get("text") or "")
            body_run.italic = True
        elif node_type == "paragraph":
            paragraph = document.add_paragraph()
            for run_node in node.get("content") or []:
                if run_node.get("type") != "text":
                    continue
                marks = {mark.get("type") for mark in (run_node.get("marks") or [])}
                run = paragraph.add_run(run_node.get("text", ""))
                run.bold = "bold" in marks
                run.italic = "italic" in marks

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def load_chapter_manuscript_for_export(cursor, chapter_id):
    """Carrega o capítulo e seu manuscrito para exportação (ADR-015), com o mesmo gate humano
    (reviewed/final) e o mesmo fallback de composição do GET /chapters/<id>/manuscript.
    Retorna (chapter, manuscript_content) ou (None, erro_response)."""
    cursor.execute("""
        SELECT ch.*, bp.title AS book_project_title
        FROM editorial.chapters ch
        JOIN editorial.book_projects bp ON bp.id = ch.book_project_id
        WHERE ch.id = %s
    """, (chapter_id,))
    chapter = cursor.fetchone()
    if not chapter:
        return None, (jsonify({"error": "Chapter not found"}), 404)
    if chapter["status"] not in ("reviewed", "final"):
        return None, (jsonify({"error": "Chapter must be reviewed or final to export", "status": chapter["status"]}), 400)

    manuscript_content = chapter["manuscript_content"] or compose_manuscript_from_sources(cursor, chapter_id)
    return (chapter, manuscript_content), None


@app.get("/chapters/<chapter_id>/export.docx")
@require_api_key
def export_chapter_docx(chapter_id):
    """ADR-015 — exportação somente-leitura do manuscrito já revisado. Não promove o status
    do capítulo; pode ser chamado quantas vezes o curador quiser enquanto ainda lapida o texto."""
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            loaded, error = load_chapter_manuscript_for_export(cursor, chapter_id)
            if error:
                return error
            chapter, manuscript_content = loaded

    buffer = build_chapter_docx(chapter["book_project_title"], chapter["title"], manuscript_content)
    ascii_name = re.sub(r"[^A-Za-z0-9_\-. ]", "_", chapter["title"]).strip() or "capitulo"
    utf8_name = quote(f"{chapter['title']}.docx")
    return Response(
        buffer.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{ascii_name}.docx\"; filename*=UTF-8''{utf8_name}"},
    )


@app.get("/chapters/<chapter_id>/export/manifest")
@require_api_key
def export_chapter_manifest(chapter_id):
    """ADR-015 §5 — rastreabilidade: lista segment_id/título de cada bloco literal exportado.
    Reconstituível a qualquer momento a partir do manuscript_content, não é persistido à parte."""
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            loaded, error = load_chapter_manuscript_for_export(cursor, chapter_id)
            if error:
                return error
            _chapter, manuscript_content = loaded

    segments = []

    def walk(node):
        if not isinstance(node, dict):
            return
        if node.get("type") == "literalSegment":
            attrs = node.get("attrs") or {}
            if attrs.get("segmentId"):
                segments.append({"segment_id": attrs["segmentId"], "title": attrs.get("title")})
        for child in node.get("content") or []:
            walk(child)

    walk(manuscript_content)
    return jsonify({"chapter_id": chapter_id, "segments": segments})


@app.post("/chapters/<chapter_id>/approve")
@require_api_key
def approve_chapter(chapter_id):
    """Gate de aprovação humana (ADR-013 §3.4): só um humano avança um capítulo de
    'draft' para 'assembled'. Nunca automatizar esta chamada."""
    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapters WHERE id = %s", (chapter_id,))
            chapter = cursor.fetchone()
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404
            if chapter["status"] != "draft":
                return jsonify({"error": "Chapter is not in draft status", "status": chapter["status"]}), 400

            cursor.execute("SELECT count(*) AS n FROM editorial.chapter_sources WHERE chapter_id = %s", (chapter_id,))
            if cursor.fetchone()["n"] == 0:
                return jsonify({"error": "Chapter has no sources to approve"}), 400

            cursor.execute("""
                UPDATE editorial.chapters SET status = 'assembled', updated_at = NOW()
                WHERE id = %s
            """, (chapter_id,))
            result = build_chapter_detail(cursor, chapter_id)
    return jsonify(result)


@app.get("/chapters/<chapter_id>/consolidation-check")
@require_api_key
def consolidation_check(chapter_id):
    """Checklist de consolidação (ADR-014 §2) — sempre um relatório informativo,
    nunca bloqueia nada sozinho. A decisão de avançar para 'reviewed' é sempre humana."""
    paraphrase_threshold = float(request.args.get("paraphrase_threshold", 0.85))

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            chapter = build_chapter_detail(cursor, chapter_id)
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404

            sources = chapter["sources"]
            issues = []

            # 1. a mesma fonte usada mais de uma vez dentro do proprio capitulo
            seen = {}
            for source in sources:
                if source["segment_id"]:
                    key = ("segment", source["segment_id"])
                elif source["knowledge_card_id"]:
                    key = ("card", source["knowledge_card_id"])
                else:
                    continue
                seen.setdefault(key, []).append(source["id"])
            for (kind, item_id), ids in seen.items():
                if len(ids) > 1:
                    issues.append({
                        "type": "duplicate_source_in_chapter",
                        "detail": f"{kind} {item_id} aparece {len(ids)} vezes neste capítulo",
                        "chapter_source_ids": [str(i) for i in ids],
                    })

            # 2. terminologia nao-canonica em texto de transicao escrito manualmente
            cursor.execute("""
                SELECT canonical_name, aliases FROM editorial.concepts
                WHERE jsonb_array_length(aliases) > 0
            """)
            concepts_with_aliases = cursor.fetchall()
            for source in sources:
                if source["inclusion_type"] == "transition_context" and source["content"]:
                    text_lower = source["content"].lower()
                    for concept in concepts_with_aliases:
                        canonical_lower = concept["canonical_name"].lower()
                        for alias in concept["aliases"]:
                            if alias.lower() in text_lower and canonical_lower not in text_lower:
                                issues.append({
                                    "type": "non_canonical_terminology",
                                    "chapter_source_id": str(source["id"]),
                                    "alias_found": alias,
                                    "canonical_name": concept["canonical_name"],
                                })

            # 3. possivel parafrase de bloco literal adjacente
            for index, source in enumerate(sources):
                if source["inclusion_type"] != "transition_context" or not source["content"]:
                    continue
                embedding = generate_embedding(source["content"])
                if embedding is None:
                    continue
                neighbors = [sources[index - 1] if index > 0 else None,
                             sources[index + 1] if index + 1 < len(sources) else None]
                for neighbor in neighbors:
                    if not neighbor or neighbor["inclusion_type"] != "literal_segment" or not neighbor["segment_id"]:
                        continue
                    cursor.execute("""
                        SELECT 1 - (embedding <=> %s::vector) AS similarity
                        FROM editorial.content_segments
                        WHERE id = %s AND embedding IS NOT NULL
                    """, (embedding, neighbor["segment_id"]))
                    row = cursor.fetchone()
                    if row and row["similarity"] >= paraphrase_threshold:
                        issues.append({
                            "type": "possible_paraphrase",
                            "chapter_source_id": str(source["id"]),
                            "adjacent_segment_id": str(neighbor["segment_id"]),
                            "similarity": row["similarity"],
                        })

    return jsonify({
        "chapter_id": chapter_id,
        "issues": issues,
        "note": (
            "Canalização cortada no meio entre dois chapter_sources não é verificado "
            "aqui porque é estruturalmente impossível: cada linha de chapter_sources "
            "referencia um segmento inteiro, nunca um trecho parcial."
        ),
    })


@app.get("/book-projects/<book_project_id>/duplicate-report")
@require_api_key
def duplicate_report(book_project_id):
    """Detecção de duplicidade entre capítulos do mesmo projeto (ADR-014 §1).
    Só sinaliza — nunca remove ou realoca nada automaticamente."""
    threshold = float(request.args.get("threshold", 0.90))

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT id FROM editorial.book_projects WHERE id = %s", (book_project_id,))
            if not cursor.fetchone():
                return jsonify({"error": "Book project not found"}), 404

            cursor.execute("""
                WITH proj_sources AS (
                    SELECT
                        cs.id, cs.chapter_id, cs.segment_id, cs.knowledge_card_id,
                        COALESCE(seg.embedding, card.embedding) AS embedding,
                        COALESCE(seg.title, card.title) AS title
                    FROM editorial.chapter_sources cs
                    JOIN editorial.chapters ch ON ch.id = cs.chapter_id
                    LEFT JOIN editorial.content_segments seg ON seg.id = cs.segment_id
                    LEFT JOIN editorial.knowledge_cards card ON card.id = cs.knowledge_card_id
                    WHERE ch.book_project_id = %(book_project_id)s
                )
                SELECT
                    a.chapter_id AS chapter_a_id, b.chapter_id AS chapter_b_id,
                    a.id AS source_a_id, b.id AS source_b_id,
                    a.title AS title_a, b.title AS title_b,
                    1 - (a.embedding <=> b.embedding) AS similarity
                FROM proj_sources a
                JOIN proj_sources b ON a.id < b.id AND a.chapter_id <> b.chapter_id
                WHERE a.embedding IS NOT NULL AND b.embedding IS NOT NULL
                  AND 1 - (a.embedding <=> b.embedding) >= %(threshold)s
                ORDER BY similarity DESC
            """, {"book_project_id": book_project_id, "threshold": threshold})
            conflicts = cursor.fetchall()

    return jsonify({"book_project_id": book_project_id, "threshold": threshold, "conflicts": conflicts})


@app.post("/chapters/<chapter_id>/review")
@require_api_key
def review_chapter(chapter_id):
    """Gate de aprovação humana (ADR-014 §3): só avança de 'assembled' para 'reviewed'
    com reviewed_by preenchido. Nunca deve ser chamado por um processo automático."""
    payload = request.get_json(silent=True) or {}
    reviewed_by = payload.get("reviewed_by")
    if not reviewed_by:
        return jsonify({"error": "Missing required fields", "fields": ["reviewed_by"]}), 400

    with get_connection() as connection:
        with connection.cursor(cursor_factory=RealDictCursor) as cursor:
            cursor.execute("SELECT * FROM editorial.chapters WHERE id = %s", (chapter_id,))
            chapter = cursor.fetchone()
            if not chapter:
                return jsonify({"error": "Chapter not found"}), 404
            if chapter["status"] != "assembled":
                return jsonify({
                    "error": "Chapter must be 'assembled' before it can be reviewed",
                    "status": chapter["status"],
                }), 400

            cursor.execute("""
                UPDATE editorial.chapters
                SET status = 'reviewed', reviewed_by = %s, reviewed_at = NOW(), updated_at = NOW()
                WHERE id = %s
            """, (reviewed_by, chapter_id))
            result = build_chapter_detail(cursor, chapter_id)
    return jsonify(result)


@app.errorhandler(Exception)
def handle_unexpected_error(error):
    app.logger.exception("Unexpected error")
    return jsonify({"error": "Internal server error", "message": str(error)}), 500
